import sys
from pathlib import Path
sys.path.append('/app')

from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from fastapi.responses import JSONResponse, FileResponse
from sqlalchemy.orm import Session
import logging
import json
import uuid
from datetime import datetime
import time

logger = logging.getLogger("uvicorn.error")

from backend.fastapi.user_contract_parser import UserContractParser
from backend.shared.database import init_db, get_db, ContractDocument, ClassificationResult
from backend.classification_agent.agent import classify_contract_task

app = FastAPI()


# 시작 시 DB 초기화
@app.on_event("startup")
async def startup_event():
    """애플리케이션 시작 시 실행"""
    logger.info("데이터베이스 초기화 중...")
    init_db()
    logger.info("데이터베이스 초기화 완료")
    
    # 지식베이스 상태 확인
    try:
        from backend.shared.services import get_knowledge_base_loader
        loader = get_knowledge_base_loader()
        status = loader.verify_knowledge_base()
        
        logger.info(f"지식베이스 상태: {status['status']}")
        logger.info(f"사용 가능한 계약 유형: {status['available_types']}")
        
        if status['missing_types']:
            logger.warning(f"누락된 계약 유형: {status['missing_types']}")
            logger.warning("ingestion CLI를 실행하여 지식베이스를 구축하세요.")
    except Exception as e:
        logger.error(f"지식베이스 상태 확인 실패: {e}")


@app.get("/")
async def root():
    return {"message": "FastAPI 서버 실행 중"}


@app.get("/api/knowledge-base/status")
async def knowledge_base_status():
    """
    지식베이스 상태 확인
    
    Returns:
        {
            "status": "ok" | "incomplete" | "missing",
            "available_types": [...],
            "missing_types": [...],
            "details": {...}
        }
    """
    try:
        from backend.shared.services import get_knowledge_base_loader
        
        loader = get_knowledge_base_loader()
        status = loader.verify_knowledge_base()
        
        return status
        
    except Exception as e:
        logger.exception(f"지식베이스 상태 확인 중 오류: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _temp_file_path(filename: str) -> Path:
    """임시 파일 경로 생성"""
    base = Path("/tmp/uploads")
    base.mkdir(parents=True, exist_ok=True)
    return base / filename


@app.post("/upload")
async def upload_file(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """
    사용자 계약서 DOCX 업로드 및 전체 처리 (파싱 → 청킹 → 임베딩 → 분류)
    
    Args:
        file: 업로드된 DOCX 파일
        db: 데이터베이스 세션
        
    Returns:
        {
            "success": bool,
            "filename": str,
            "contract_id": str,
            "chunks_count": int,
            "metadata": dict,
            "message": str
        }
    """
    try:
        filename = Path(file.filename).name
        
        # DOCX 파일만 허용
        if not filename.lower().endswith('.docx'):
            raise HTTPException(status_code=400, detail="DOCX 파일만 허용됩니다.")

        # contract_id 생성
        contract_id = f"contract_{uuid.uuid4().hex[:12]}"
        
        # 파일 저장 디렉토리
        save_dir = Path("/app/data/source_documents")
        save_dir.mkdir(parents=True, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        saved_filename = f"user_contract_{timestamp}.docx"
        saved_path = save_dir / saved_filename
        
        # 파일 저장
        content = await file.read()
        with open(saved_path, 'wb') as f:
            f.write(content)
        
        logger.info(f"파일 저장 완료: {saved_path}")

        # UserContractProcessor 초기화
        from backend.fastapi.user_contract_parser import UserContractParser
        from backend.shared.services.user_contract_chunker import UserContractChunker
        from backend.shared.services.user_contract_embedder import UserContractEmbedder
        from backend.shared.services.user_contract_processor import UserContractProcessor
        from backend.shared.services import get_embedding_service
        
        parser = UserContractParser()
        chunker = UserContractChunker()
        embedding_service = get_embedding_service()
        embedder = UserContractEmbedder(embedding_service)
        processor = UserContractProcessor(parser, chunker, embedder)
        
        # 출력 디렉토리
        output_dir = Path("/app/data/user_contracts") / contract_id
        output_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"처리 시작: {contract_id}")
        
        # 전체 파이프라인 실행 (파싱 → 청킹 → 임베딩)
        process_result = processor.process_contract(
            docx_path=saved_path,
            output_dir=output_dir,
            contract_id=contract_id
        )
        
        if not process_result["success"]:
            raise HTTPException(
                status_code=500,
                detail=f"처리 실패: {process_result.get('error', 'Unknown error')}"
            )
        
        logger.info(f"처리 완료: {contract_id}, {process_result['metadata']['total_chunks']}개 청크")
        
        # DB에 저장
        contract_doc = ContractDocument(
            contract_id=contract_id,
            filename=filename,
            file_path=str(saved_path),
            parsed_data=None,  # chunks.json 경로로 대체
            parsed_metadata={
                "chunks_path": process_result["chunks_path"],
                "embeddings_path": process_result["embeddings_path"],
                **process_result["metadata"]
            },
            status="processed"
        )
        db.add(contract_doc)
        db.commit()
        
        logger.info(f"DB 저장 완료: {contract_id}")

        # Celery를 통해 분류 작업을 큐에 전송
        try:
            task = classify_contract_task.delay(contract_id)
            logger.info(f"분류 작업 큐에 전송: {contract_id}, Task ID: {task.id}")

            # 계약서 상태를 classifying으로 업데이트
            contract_doc.status = "classifying"
            db.commit()

        except Exception as e:
            logger.error(f"분류 작업 큐 전송 실패: {e}")

        return JSONResponse(
            content={
                "success": True,
                "filename": filename,
                "contract_id": contract_id,
                "chunks_count": process_result["metadata"]["total_chunks"],
                "parsed_metadata": process_result["metadata"],
                "structured_data": process_result.get("structured_data", {}),
                "metadata": process_result["metadata"],
                "message": "처리 완료. 분류 작업이 백그라운드에서 진행 중입니다."
            },
            media_type="application/json; charset=utf-8"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"업로드 처리 중 오류: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/upload-simple")
async def upload_file_simple(file: UploadFile = File(...)):
    """
    간단한 파일 업로드 및 파싱 (PDF, DOCX, TXT 지원, DB 저장 없음)
    
    Args:
        file: 업로드된 파일
        
    Returns:
        {
            "success": bool,
            "filename": str,
            "file_type": str,
            "size_bytes": int,
            "text": str (또는 파싱 결과)
        }
    """
    try:
        filename = Path(file.filename).name
        file_ext = filename.lower().split('.')[-1]
        
        # 지원 형식 확인
        if file_ext not in ['pdf', 'docx', 'txt']:
            raise HTTPException(
                status_code=400, 
                detail="지원하지 않는 파일 형식입니다. (지원: PDF, DOCX, TXT)"
            )

        # 임시 파일 저장
        temp_path = _temp_file_path(filename)
        content = await file.read()
        
        # 파일 저장
        with open(temp_path, 'wb') as f:
            f.write(content)

        # 파일 형식에 따라 처리
        result = {}
        
        if file_ext == 'pdf':
            # PyMuPDF 파싱
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(str(temp_path))
                text_parts = []
                for page in doc:
                    text_parts.append(page.get_text())
                text = "\n".join(text_parts)
                doc.close()
                result["text"] = text
                result["pages"] = len(doc)
            except ImportError:
                raise HTTPException(
                    status_code=500,
                    detail="PDF 처리를 위해 PyMuPDF가 필요합니다."
                )
        
        elif file_ext == 'docx':
            # DOCX 텍스트 추출
            try:
                from docx import Document
                doc = Document(str(temp_path))
                text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                result["text"] = text
                result["paragraphs"] = len([p for p in doc.paragraphs if p.text.strip()])
            except ImportError:
                raise HTTPException(
                    status_code=500, 
                    detail="DOCX 처리를 위해 python-docx가 필요합니다."
                )
        
        elif file_ext == 'txt':
            # TXT 파일 읽기 (여러 인코딩 시도)
            encodings = ['utf-8', 'cp949', 'euc-kr', 'utf-8-sig']
            text = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(temp_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise HTTPException(
                    status_code=400,
                    detail=f"파일을 읽을 수 없습니다. 시도한 인코딩: {', '.join(encodings)}"
                )
            
            result["text"] = text
            result["lines"] = len(text.split('\n'))
            result["encoding"] = used_encoding

        # 임시 파일 삭제
        try:
            temp_path.unlink()
        except Exception as e:
            logger.warning(f"임시 파일 삭제 실패: {e}")

        return {
            "success": True,
            "filename": filename,
            "file_type": file_ext,
            "size_bytes": len(content),
            **result
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"파일 처리 오류: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/verify")
async def verify_contract(file: UploadFile = File(...), contract_type: str = "provide"):
    """
    계약서 업로드 및 자동 검증
    
    Args:
        file: 업로드된 계약서 파일 (TXT, DOCX, PDF)
        contract_type: 계약서 유형 (provide, create, process, brokerage_provider, brokerage_user)
    
    Returns:
        검증 결과 요약 및 리포트 경로
    """
    try:
        start_time = time.time()
        
        filename = Path(file.filename).name
        file_ext = filename.lower().split('.')[-1]
        
        # 지원 형식 확인
        if file_ext not in ['txt', 'docx', 'pdf']:
            raise HTTPException(
                status_code=400,
                detail="지원하지 않는 파일 형식입니다. (지원: TXT, DOCX, PDF)"
            )
        
        logger.info(f"[Verify] 파일 업로드: {filename}")
        
        # 파일 저장
        save_dir = Path("/app/data/source_documents")
        save_dir.mkdir(parents=True, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        saved_filename = f"user_contract_{timestamp}.{file_ext}"
        saved_path = save_dir / saved_filename
        
        content = await file.read()
        with open(saved_path, 'wb') as f:
            f.write(content)
        
        logger.info(f"[Verify] 파일 저장: {saved_path}")
        
        # 검증 실행
        logger.info(f"[Verify] 검증 시작... (계약 유형: {contract_type})")
        result = _run_verification_api(str(saved_path), contract_type)
        
        execution_time = time.time() - start_time
        
        logger.info(f"[Verify] 검증 완료 ({execution_time:.2f}s)")
        
        return {
            "success": True,
            "filename": filename,
            "report_id": result["report_id"],
            "report_path": result["report_path"],
            "verification_summary": {
                "total_standard_clauses": result["total_standard_clauses"],
                "total_user_clauses": result["total_user_clauses"],
                "matched_clauses": result["matched_clauses"],
                "missing_clauses": result["missing_count"],
                "compliance_rate": result["verification_rate"]
            },
            "execution_time": execution_time
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[Verify] 검증 오류: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"검증 실패: {str(e)}")


def _run_verification_api(user_contract_path: str, contract_type: str = "provide"):
    """
    API용 검증 실행 함수 (ingestion/ingest.py의 _run_verification 로직 재사용)
    
    Args:
        user_contract_path: 사용자 계약서 파일 경로
        contract_type: 계약서 유형 (provide, create, process, brokerage_provider, brokerage_user)
    
    Returns:
        검증 결과 딕셔너리
    """
    from backend.clause_verification.node_1_clause_matching.data_loader import ContractDataLoader
    from backend.clause_verification.node_1_clause_matching.verification_engine import ContractVerificationEngine
    from backend.clause_verification.node_1_clause_matching.hybrid_search import HybridSearchEngine
    from backend.clause_verification.node_1_clause_matching.llm_verification import LLMVerificationService
    from backend.shared.services import get_embedding_service
    
    # 계약 유형별 표준 계약서 매핑
    contract_type_mapping = {
        "provide": "provide_std_contract_chunks.json",
        "create": "create_std_contract_chunks.json",
        "process": "process_std_contract_chunks.json",
        "brokerage_provider": "brokerage_provider_std_contract_chunks.json",
        "brokerage_user": "brokerage_user_std_contract_chunks.json"
    }
    
    # 표준 계약서 경로
    standard_filename = contract_type_mapping.get(contract_type, "provide_std_contract_chunks.json")
    standard_contract_path = f"/app/data/chunked_documents/{standard_filename}"
    
    if not Path(standard_contract_path).exists():
        raise FileNotFoundError(f"표준 계약서를 찾을 수 없습니다: {standard_contract_path}")
    
    if not Path(user_contract_path).exists():
        raise FileNotFoundError(f"사용자 계약서를 찾을 수 없습니다: {user_contract_path}")
    
    logger.info("[Verify] 1단계: 데이터 로드")
    
    # 데이터 로더 초기화
    loader = ContractDataLoader()
    
    # 표준 계약서 로드 (KnowledgeBaseLoader 사용)
    standard_clauses = loader.load_standard_contract(
        contract_type=contract_type,
        use_knowledge_base=True
    )
    logger.info(f"[Verify] 표준 계약서 로드: {len(standard_clauses)}개 조문 (유형: {contract_type})")
    
    # 사용자 계약서 로드 (파일 형식에 따라 처리)
    file_ext = Path(user_contract_path).suffix.lower()
    
    if file_ext == '.txt':
        # TXT 파일: 여러 인코딩 시도
        encodings = ['utf-8', 'cp949', 'euc-kr', 'utf-8-sig', 'latin-1']
        user_text = None
        for encoding in encodings:
            try:
                with open(user_contract_path, 'r', encoding=encoding) as f:
                    user_text = f.read()
                logger.info(f"[Verify] 파일 인코딩 감지: {encoding}")
                break
            except (UnicodeDecodeError, LookupError):
                continue
        
        if user_text is None:
            raise ValueError(f"파일을 읽을 수 없습니다. 시도한 인코딩: {', '.join(encodings)}")
    
    elif file_ext == '.docx':
        # DOCX 파일: python-docx 사용
        try:
            from docx import Document
            doc = Document(user_contract_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            user_text = "\n".join(paragraphs)
            logger.info(f"[Verify] DOCX 파일 처리: {len(paragraphs)}개 문단")
        except ImportError:
            raise ImportError("DOCX 처리를 위해 python-docx가 필요합니다")
    
    elif file_ext == '.pdf':
        # PDF 파일: PyMuPDF 사용
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(user_contract_path)
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            user_text = "\n".join(text_parts)
            doc.close()
            logger.info(f"[Verify] PDF 파일 처리: {len(doc)}페이지")
        except ImportError:
            raise ImportError("PDF 처리를 위해 PyMuPDF가 필요합니다")
    
    else:
        raise ValueError(f"지원하지 않는 파일 형식: {file_ext}")
    
    if not user_text or not user_text.strip():
        raise ValueError("파일에서 텍스트를 추출할 수 없습니다")
    
    user_clauses = loader.load_user_contract_chunked(user_text)
    logger.info(f"[Verify] 사용자 계약서 로드: {len(user_clauses)}개 청크")
    
    logger.info("[Verify] 2단계: 검증 엔진 초기화")
    
    # 서비스 초기화
    embedding_service = get_embedding_service()
    hybrid_search = HybridSearchEngine()  # 역방향 검증에서는 사용 안 함
    llm_verification = LLMVerificationService()
    
    # 검증 엔진 초기화
    engine = ContractVerificationEngine(
        embedding_service=embedding_service,
        hybrid_search=hybrid_search,
        llm_verification=llm_verification
    )
    
    logger.info("[Verify] 3단계: 계약서 검증 수행")
    
    # 검증 수행
    result = engine.verify_contract_reverse(
        standard_clauses=standard_clauses,
        user_clauses=user_clauses,
        top_k_candidates=10,
        top_k_titles=5,
        min_confidence=0.5
    )
    
    logger.info(f"[Verify] 검증 완료 - 매칭: {result.matched_clauses}/{result.total_user_clauses}")
    
    logger.info("[Verify] 4단계: 보고서 생성")
    
    # 보고서 생성
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_path = Path(f"/app/data/reports/verification_report_{timestamp}.txt")
    report_path.parent.mkdir(parents=True, exist_ok=True)
    
    # 간단한 보고서 생성
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write("="*100 + "\n")
        f.write("계약서 검증 보고서\n")
        f.write("="*100 + "\n\n")
        f.write(f"생성 시각: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        
        f.write("📊 검증 결과 요약\n")
        f.write("-" * 50 + "\n")
        f.write(f"표준 조문 수: {result.total_standard_clauses}개\n")
        f.write(f"사용자 청크 수: {result.total_user_clauses}개\n")
        f.write(f"매칭된 청크: {result.matched_clauses}개\n")
        f.write(f"누락된 조문: {result.missing_count}개\n")
        f.write(f"검증 완료율: {result.verification_rate:.1f}%\n\n")
        
        # 매칭된 조문
        if result.match_results:
            f.write("✅ 매칭된 조문\n")
            f.write("-" * 50 + "\n")
            for match in [m for m in result.match_results if m.is_matched]:
                f.write(f"[{match.standard_clause.id}] {match.standard_clause.title}\n")
                f.write(f"  ← {match.matched_clause.title}\n")
                if match.llm_decision:
                    f.write(f"  신뢰도: {match.llm_decision.confidence:.0%}\n")
                f.write("\n")
        
        # 누락된 조문
        if result.missing_clauses:
            f.write("\n❌ 누락된 조문\n")
            f.write("-" * 50 + "\n")
            for clause in result.missing_clauses:
                f.write(f"[{clause.id}] {clause.title}\n")
                f.write(f"  {clause.text[:100]}...\n\n")
    
    logger.info(f"[Verify] 보고서 저장: {report_path}")
    
    return {
        "report_id": timestamp,
        "report_path": str(report_path),
        "total_standard_clauses": result.total_standard_clauses,
        "total_user_clauses": result.total_user_clauses,
        "matched_clauses": result.matched_clauses,
        "missing_count": result.missing_count,
        "verification_rate": result.verification_rate
    }


@app.get("/report/{report_id}")
async def download_report(report_id: str):
    """
    검증 리포트 다운로드
    
    Args:
        report_id: 리포트 ID (타임스탬프)
    
    Returns:
        FileResponse: 리포트 파일
    """
    try:
        report_path = Path(f"/app/data/reports/verification_report_{report_id}.txt")
        
        if not report_path.exists():
            raise HTTPException(status_code=404, detail="리포트를 찾을 수 없습니다")
        
        return FileResponse(
            path=str(report_path),
            media_type="text/plain",
            filename=f"verification_report_{report_id}.txt"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"리포트 다운로드 오류: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/classification/{contract_id}/start")
async def start_classification(contract_id: str, db: Session = Depends(get_db)):
    """
    계약서 분류 시작 (수동 트리거)

    Args:
        contract_id: 계약서 ID
        db: 데이터베이스 세션

    Returns:
        {
            "success": bool,
            "contract_id": str,
            "task_id": str,
            "message": str
        }
    """
    try:
        # 계약서 조회
        contract = db.query(ContractDocument).filter(
            ContractDocument.contract_id == contract_id
        ).first()

        if not contract:
            raise HTTPException(status_code=404, detail="계약서를 찾을 수 없습니다")

        if not contract.parsed_data:
            raise HTTPException(status_code=400, detail="파싱된 데이터가 없습니다")

        # Celery Task 큐에 전송
        task = classify_contract_task.delay(contract_id)

        # 계약서 상태 업데이트
        contract.status = "classifying"
        db.commit()

        logger.info(f"분류 작업 큐에 전송: {contract_id}, Task ID: {task.id}")

        return {
            "success": True,
            "contract_id": contract_id,
            "task_id": task.id,
            "message": "분류 작업이 시작되었습니다. /api/classification/{contract_id}로 결과를 조회하세요."
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"분류 시작 중 오류: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/classification/{contract_id}")
async def get_classification(contract_id: str, db: Session = Depends(get_db)):
    """
    분류 결과 조회

    Args:
        contract_id: 계약서 ID
        db: 데이터베이스 세션

    Returns:
        분류 결과
    """
    try:
        classification = db.query(ClassificationResult).filter(
            ClassificationResult.contract_id == contract_id
        ).first()

        if not classification:
            raise HTTPException(status_code=404, detail="분류 결과를 찾을 수 없습니다")

        return {
            "contract_id": classification.contract_id,
            "predicted_type": classification.predicted_type,
            "confidence": classification.confidence,
            "scores": classification.scores,
            "confirmed_type": classification.confirmed_type,
            "user_override": classification.user_override
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"분류 조회 중 오류: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/classification/{contract_id}/confirm")
async def confirm_classification(
    contract_id: str,
    confirmed_type: str,
    db: Session = Depends(get_db)
):
    """
    사용자가 분류 유형 확인/수정

    Args:
        contract_id: 계약서 ID
        confirmed_type: 사용자가 확인한 유형
        db: 데이터베이스 세션

    Returns:
        {
            "success": bool,
            "contract_id": str,
            "confirmed_type": str
        }
    """
    try:
        classification = db.query(ClassificationResult).filter(
            ClassificationResult.contract_id == contract_id
        ).first()

        if not classification:
            raise HTTPException(status_code=404, detail="분류 결과를 찾을 수 없습니다")

        # 사용자가 변경한 경우 기록
        if confirmed_type != classification.predicted_type:
            classification.user_override = confirmed_type

        classification.confirmed_type = confirmed_type

        # 계약서 상태 업데이트
        contract = db.query(ContractDocument).filter(
            ContractDocument.contract_id == contract_id
        ).first()

        if contract:
            contract.status = "classified_confirmed"

        db.commit()

        logger.info(f"분류 확인: {contract_id} -> {confirmed_type}")

        return {
            "success": True,
            "contract_id": contract_id,
            "confirmed_type": confirmed_type
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"분류 확인 중 오류: {e}")
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)


@app.post("/api/consistency/{contract_id}/start")
async def start_consistency_verification(contract_id: str, db: Session = Depends(get_db)):
    """
    정합성 검증 시작
    
    Args:
        contract_id: 계약서 ID
        db: 데이터베이스 세션
        
    Returns:
        {
            "success": bool,
            "contract_id": str,
            "task_id": str
        }
    """
    try:
        # 계약서 확인
        contract = db.query(ContractDocument).filter(
            ContractDocument.contract_id == contract_id
        ).first()
        
        if not contract:
            raise HTTPException(status_code=404, detail="계약서를 찾을 수 없습니다")
        
        # 분류 결과 확인
        classification = db.query(ClassificationResult).filter(
            ClassificationResult.contract_id == contract_id
        ).first()
        
        if not classification or not classification.confirmed_type:
            raise HTTPException(status_code=400, detail="분류가 완료되지 않았습니다")
        
        # 검증 작업 큐에 전송
        from backend.consistency_agent.agent import verify_contract_task
        task = verify_contract_task.delay(contract_id)
        
        # 계약서 상태 업데이트
        contract.status = "verifying"
        db.commit()
        
        logger.info(f"정합성 검증 시작: {contract_id}, Task ID: {task.id}")
        
        return {
            "success": True,
            "contract_id": contract_id,
            "task_id": task.id
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"정합성 검증 시작 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/consistency/{contract_id}")
async def get_consistency_result(contract_id: str, db: Session = Depends(get_db)):
    """
    정합성 검증 결과 조회
    
    Args:
        contract_id: 계약서 ID
        db: 데이터베이스 세션
        
    Returns:
        검증 결과 또는 진행 상태
    """
    try:
        contract = db.query(ContractDocument).filter(
            ContractDocument.contract_id == contract_id
        ).first()
        
        if not contract:
            raise HTTPException(status_code=404, detail="계약서를 찾을 수 없습니다")
        
        return {
            "contract_id": contract_id,
            "status": contract.status,
            "result": contract.verification_result if hasattr(contract, 'verification_result') else None
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"정합성 검증 결과 조회 실패: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/consistency/{contract_id}/report")
async def get_verification_report(contract_id: str, db: Session = Depends(get_db)):
    """
    검증 결과를 텍스트 리포트로 반환
    
    Args:
        contract_id: 계약서 ID
        db: 데이터베이스 세션
        
    Returns:
        텍스트 리포트
    """
    try:
        from backend.consistency_agent.node_1_clause_matching.data_loader import ContractDataLoader
        from backend.consistency_agent.node_1_clause_matching.verifier import ContractVerificationEngine
        from backend.consistency_agent.node_1_clause_matching.hybrid_search import HybridSearchEngine
        from backend.consistency_agent.node_1_clause_matching.llm_verification import LLMVerificationService
        from backend.shared.services.embedding_service import EmbeddingService
        from backend.shared.database import ClassificationResult
        import json
        
        # 계약서 조회
        contract = db.query(ContractDocument).filter(
            ContractDocument.contract_id == contract_id
        ).first()
        
        if not contract:
            raise HTTPException(status_code=404, detail="계약서를 찾을 수 없습니다")
        
        if contract.status != "verified":
            raise HTTPException(status_code=400, detail=f"검증이 완료되지 않았습니다. 현재 상태: {contract.status}")
        
        # 분류 결과 조회
        classification = db.query(ClassificationResult).filter(
            ClassificationResult.contract_id == contract_id
        ).first()
        
        contract_type = classification.confirmed_type or classification.predicted_type
        chunks_path = contract.parsed_metadata.get("chunks_path")
        
        # 검증 재실행
        loader = ContractDataLoader()
        embedding_service = EmbeddingService()
        hybrid_search = HybridSearchEngine()
        llm_verification = LLMVerificationService()
        verifier = ContractVerificationEngine(
            embedding_service=embedding_service,
            hybrid_search=hybrid_search,
            llm_verification=llm_verification,
            data_loader=loader
        )
        
        standard_clauses = loader.load_standard_contract(
            contract_type=contract_type,
            use_knowledge_base=True
        )
        
        with open(chunks_path, 'r', encoding='utf-8') as f:
            user_chunks = json.load(f)
        
        from backend.consistency_agent.node_1_clause_matching.models import ClauseData
        user_clauses = [
            ClauseData(
                id=chunk['id'],
                title=chunk.get('title', ''),
                subtitle=None,
                type=chunk.get('unit_type', 'article'),
                text=chunk.get('text_raw', ''),
                text_norm=chunk.get('text_norm', ''),
                breadcrumb=chunk.get('title', ''),
                embedding=chunk.get('embedding')
            )
            for chunk in user_chunks
        ]
        
        result = verifier.verify_contract_reverse(
            standard_clauses=standard_clauses,
            user_clauses=user_clauses
        )
        
        # 간단한 텍스트 리포트 생성
        report_lines = []
        report_lines.append("=" * 80)
        report_lines.append("계약서 검증 결과")
        report_lines.append("=" * 80)
        report_lines.append(f"\n계약서 ID: {contract_id}")
        report_lines.append(f"계약 유형: {contract_type}")
        report_lines.append(f"\n검증 결과:")
        report_lines.append(f"  - 표준 조문: {result.total_standard_clauses}개")
        report_lines.append(f"  - 사용자 청크: {result.total_user_clauses}개")
        report_lines.append(f"  - 매칭됨: {result.matched_clauses}개")
        report_lines.append(f"  - 누락됨: {result.missing_count}개")
        report_lines.append(f"  - 검증율: {result.verification_rate:.1f}%")
        report_lines.append(f"\n{'='*80}\n")
        
        # 매칭된 조문
        report_lines.append("✅ 매칭된 조문:")
        for match in [m for m in result.match_results if m.is_matched]:
            report_lines.append(f"\n  [{match.standard_clause.id}] {match.standard_clause.title}")
            report_lines.append(f"    ← {match.matched_clause.title}")
            if match.llm_decision:
                report_lines.append(f"    신뢰도: {match.llm_decision.confidence:.0%}")
        
        # 누락된 조문
        report_lines.append(f"\n\n{'='*80}")
        report_lines.append("❌ 누락된 조문:")
        for clause in result.missing_clauses:
            report_lines.append(f"\n  [{clause.id}] {clause.title}")
            report_lines.append(f"    {clause.text[:100]}...")
        
        report_lines.append(f"\n{'='*80}")
        
        return {"report": "\n".join(report_lines)}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"리포트 생성 실패: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))
